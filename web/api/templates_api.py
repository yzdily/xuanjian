"""
自定义报告模版管理 API。

URL 保持不变：/api/templates/*
"""
from __future__ import annotations

import json
import os
import time
from pathlib import Path

from fastapi import APIRouter, Request

from core.log import get_logger

log = get_logger("web.templates_api")

router = APIRouter()


CUSTOM_TEMPLATE_DIR = Path("data/custom_templates")
CUSTOM_TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
CUSTOM_TEMPLATE_META = CUSTOM_TEMPLATE_DIR / "_meta.json"

# 支持的上传格式
_ALLOWED_TEMPLATE_EXTS = {".pdf", ".docx", ".md", ".txt", ".html"}
_MAX_TEMPLATE_SIZE = 20 * 1024 * 1024  # 20MB
_TEMPLATE_TEXT_LIMIT = 100000


def _load_template_meta() -> list[dict]:
    """加载自定义模版元数据列表。"""
    if not CUSTOM_TEMPLATE_META.exists():
        return []
    try:
        return json.loads(CUSTOM_TEMPLATE_META.read_text(encoding="utf-8"))
    except Exception:
        return []


def _save_template_meta(meta: list[dict]) -> None:
    """保存自定义模版元数据列表（原子写）。"""
    CUSTOM_TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    tmp_path = CUSTOM_TEMPLATE_META.with_suffix(CUSTOM_TEMPLATE_META.suffix + ".tmp")
    tmp_path.write_text(
        json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    os.replace(tmp_path, CUSTOM_TEMPLATE_META)


def _extract_template_text(file_path: Path, ext: str) -> str:
    """从上传的模版文件中提取纯文本内容，供 LLM 理解模版结构。"""
    try:
        if ext in (".md", ".txt", ".html"):
            return file_path.read_text(encoding="utf-8", errors="ignore")[:_TEMPLATE_TEXT_LIMIT]
        elif ext == ".pdf":
            return _extract_pdf_text(file_path)
        elif ext == ".docx":
            return _extract_docx_text(file_path)
        elif ext == ".doc":
            return "[.doc 格式不受支持，请在 Word 中另存为 .docx 后重传]"
    except Exception as e:
        log.warning("提取模版文本失败 (%s): %s", file_path.name, e)
        return f"[文本提取失败: {e}]"
    return "[不支持的格式]"


def _extract_pdf_text(file_path: Path) -> str:
    """从 PDF 文件提取文本。"""
    try:
        import PyPDF2
        text_parts = []
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages[:50]:
                page_text = page.extract_text() or ""
                text_parts.append(page_text)
        return "\n\n".join(text_parts)[:_TEMPLATE_TEXT_LIMIT]
    except ImportError:
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages[:50]:
                    text_parts.append(page.extract_text() or "")
            return "\n\n".join(text_parts)[:_TEMPLATE_TEXT_LIMIT]
        except ImportError:
            return "[需要安装 PyPDF2 或 pdfplumber 库来解析 PDF 文件。请运行: pip install PyPDF2]"


def _extract_docx_text(file_path: Path) -> str:
    """从 Word 文件提取文本。"""
    try:
        import docx
        doc = docx.Document(str(file_path))
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                style_name = para.style.name if para.style else ""
                if "Heading" in style_name or "标题" in style_name:
                    level = ""
                    for ch in style_name:
                        if ch.isdigit():
                            level = ch
                            break
                    prefix = "#" * (int(level) if level else 1)
                    text_parts.append(f"{prefix} {para.text}")
                else:
                    text_parts.append(para.text)
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                text_parts.append("\n".join(rows))
        return "\n\n".join(text_parts)[:_TEMPLATE_TEXT_LIMIT]
    except ImportError:
        return "[需要安装 python-docx 库来解析 Word 文件。请运行: pip install python-docx]"


@router.get("/api/templates/list")
async def templates_list():
    """返回所有自定义报告模版列表。"""
    meta = _load_template_meta()
    return {"status": "ok", "templates": meta, "count": len(meta)}


@router.post("/api/templates/upload")
async def templates_upload(request: Request):
    """上传自定义报告模版文件（支持 PDF、Word、Markdown、TXT、HTML）。"""
    import uuid as _uuid
    form = await request.form()
    file = form.get("file")
    if not file:
        return {"status": "error", "message": "未提供文件"}

    filename = getattr(file, "filename", "") or "unknown"
    ext = Path(filename).suffix.lower()
    if ext == ".doc":
        return {"status": "error",
                "message": ".doc（Word 97-2003 老二进制格式）无法解析。请在 Word 中打开后 “另存为→.docx” 再上传。"}
    if ext not in _ALLOWED_TEMPLATE_EXTS:
        return {"status": "error", "message": f"不支持的文件格式: {ext}，支持: {', '.join(sorted(_ALLOWED_TEMPLATE_EXTS))}"}

    content = await file.read()
    if len(content) > _MAX_TEMPLATE_SIZE:
        return {"status": "error", "message": f"文件过大（{len(content) / 1024 / 1024:.1f}MB），最大 20MB"}

    template_id = _uuid.uuid4().hex[:12]
    safe_filename = f"{template_id}{ext}"
    save_path = CUSTOM_TEMPLATE_DIR / safe_filename
    save_path.write_bytes(content)

    extracted_text = _extract_template_text(save_path, ext)

    text_path = CUSTOM_TEMPLATE_DIR / f"{template_id}.extracted.txt"
    text_path.write_text(extracted_text, encoding="utf-8")

    name = (form.get("name") or "").strip() or Path(filename).stem
    description = (form.get("description") or "").strip()
    meta = _load_template_meta()
    meta.append({
        "id": template_id,
        "name": name,
        "description": description,
        "filename": filename,
        "saved_as": safe_filename,
        "ext": ext,
        "size": len(content),
        "enabled": True,
        "uploaded_at": time.time(),
        "text_length": len(extracted_text),
    })
    _save_template_meta(meta)

    return {
        "status": "ok",
        "template_id": template_id,
        "name": name,
        "filename": filename,
        "size": len(content),
        "text_length": len(extracted_text),
    }


@router.post("/api/templates/toggle")
async def templates_toggle(request: Request):
    """切换模版启用/禁用状态（单选语义）。"""
    body = await request.json()
    template_id = (body.get("id") or "").strip()
    enabled = bool(body.get("enabled", True))

    meta = _load_template_meta()
    target = next((m for m in meta if m["id"] == template_id), None)
    if target is None:
        return {"status": "error", "message": f"模版不存在: {template_id}"}

    if enabled:
        for item in meta:
            item["enabled"] = (item["id"] == template_id)
    else:
        target["enabled"] = False

    _save_template_meta(meta)
    return {"status": "ok", "id": template_id, "enabled": enabled}


@router.post("/api/templates/meta/update")
async def templates_meta_update(request: Request):
    """修改模版的名称 / 描述（不涉及内容）。"""
    body = await request.json()
    template_id = (body.get("id") or "").strip()
    name = (body.get("name") or "").strip()
    description = (body.get("description") or "").strip()

    if not template_id:
        return {"status": "error", "message": "缺少 id"}
    if not name:
        return {"status": "error", "message": "模版名称不能为空"}

    meta = _load_template_meta()
    item = next((m for m in meta if m["id"] == template_id), None)
    if not item:
        return {"status": "error", "message": f"模版不存在: {template_id}"}

    item["name"] = name
    item["description"] = description
    item["updated_at"] = time.time()
    _save_template_meta(meta)
    return {"status": "ok", "id": template_id, "name": name}


@router.post("/api/templates/clone")
async def templates_clone(request: Request):
    """克隆一个现有模版。"""
    import uuid as _uuid
    body = await request.json()
    template_id = (body.get("id") or "").strip()

    meta = _load_template_meta()
    src = next((m for m in meta if m.get("id") == template_id), None)
    if not src:
        return {"status": "error", "message": f"模版不存在: {template_id}"}

    src_file = CUSTOM_TEMPLATE_DIR / src["saved_as"]
    src_text = CUSTOM_TEMPLATE_DIR / f"{template_id}.extracted.txt"
    if not src_file.exists():
        return {"status": "error", "message": "原模版文件丢失"}

    new_id = _uuid.uuid4().hex[:12]
    new_ext = src.get("ext", "")
    new_saved = f"{new_id}{new_ext}"
    (CUSTOM_TEMPLATE_DIR / new_saved).write_bytes(src_file.read_bytes())
    if src_text.exists():
        (CUSTOM_TEMPLATE_DIR / f"{new_id}.extracted.txt").write_bytes(src_text.read_bytes())

    new_item = {
        **src,
        "id": new_id,
        "name": (src.get("name") or "模版") + " - 副本",
        "saved_as": new_saved,
        "enabled": False,
        "uploaded_at": time.time(),
        "updated_at": time.time(),
    }
    meta.append(new_item)
    _save_template_meta(meta)
    return {"status": "ok", "id": new_id, "name": new_item["name"]}


@router.post("/api/templates/test_render")
async def templates_test_render(request: Request):
    """使用 mock 漏洞数据 + 当前模版，让 LLM 试生成一段报告。"""
    body = await request.json()
    template_id = (body.get("id") or "").strip()
    inline_content = body.get("content")
    inline_name = (body.get("name") or "草稿模版").strip()

    template_text = ""
    template_name = ""
    if inline_content:
        template_text = str(inline_content)[:_TEMPLATE_TEXT_LIMIT]
        template_name = inline_name
    else:
        meta = _load_template_meta()
        item = next((m for m in meta if m.get("id") == template_id), None)
        if not item:
            return {"status": "error", "message": f"模版不存在: {template_id}"}
        text_path = CUSTOM_TEMPLATE_DIR / f"{template_id}.extracted.txt"
        if not text_path.exists():
            return {"status": "error", "message": "模版提取文本不存在"}
        template_text = text_path.read_text(encoding="utf-8")[:_TEMPLATE_TEXT_LIMIT]
        template_name = item.get("name", "模版")

    mock_findings = [
        {"name": "反射型 XSS", "severity": "高", "url": "https://demo.example.com/search?q=<svg/onload=alert(1)>",
         "poc": "GET /search?q=<svg/onload=alert(1)>", "impact": "任意 JS 执行、会话劫持",
         "fix": "输出端 HTML 转义 + CSP"},
        {"name": "SQL 注入", "severity": "严重", "url": "https://demo.example.com/api/user?id=1",
         "poc": "id=1 AND SLEEP(5)", "impact": "全库读取",
         "fix": "参数化查询 + 最小权限"},
        {"name": "未授权访问", "severity": "中", "url": "https://demo.example.com/admin/users",
         "poc": "直接访问无 Cookie 也返回 200", "impact": "敏感数据泄露",
         "fix": "补充鉴权中间件"},
    ]

    import time as _time, asyncio
    t0 = _time.time()

    llm_output = ""
    try:
        from core.llm.client import chat as llm_chat  # type: ignore
        prompt = (
            "你是渗透测试报告撰写助手。请严格按照下面这份『用户自定义模版』的章节结构、"
            "写作风格和排版要求，结合提供的 mock 漏洞数据，生成一份完整的报告示例。\n\n"
            f"==== 用户模版（{template_name}）====\n{template_text}\n\n"
            f"==== mock 漏洞数据 ====\n{json.dumps(mock_findings, ensure_ascii=False, indent=2)}\n\n"
            "请直接输出报告内容（Markdown）。"
        )
        if asyncio.iscoroutinefunction(llm_chat):
            llm_output = await llm_chat(prompt)
        else:
            llm_output = await asyncio.to_thread(llm_chat, prompt)
    except Exception as e:
        log.warning("test_render 调用 LLM 失败，使用本地兜底: %s", e)
        lines = [
            f"# 模版试渲染：{template_name}",
            "",
            "> ⚠️ 当前为本地兜底渲染（未调用 LLM），仅用于结构预览。真实报告生成时会按完整模版调用 LLM。",
            "",
            "## 模版片段（前 1500 字）",
            "",
            "```",
            template_text[:1500],
            "```",
            "",
            "## Mock 漏洞清单",
            "",
            "| 编号 | 漏洞 | 风险 | URL |",
            "|------|------|------|-----|",
        ]
        for i, f in enumerate(mock_findings, 1):
            lines.append(f"| {i} | {f['name']} | {f['severity']} | `{f['url']}` |")
        lines += ["", "## 漏洞详情"]
        for i, f in enumerate(mock_findings, 1):
            lines += [
                f"### {i}. {f['name']}（{f['severity']}）",
                f"- URL：`{f['url']}`",
                f"- PoC：`{f['poc']}`",
                f"- 影响：{f['impact']}",
                f"- 修复建议：{f['fix']}",
                "",
            ]
        llm_output = "\n".join(lines)

    return {
        "status": "ok",
        "template_name": template_name,
        "content": llm_output or "(LLM 未返回内容)",
        "elapsed": round(_time.time() - t0, 2),
    }


# 内置示例模版（首次进入 / 列表为空时引导用户克隆）
_BUILTIN_TEMPLATES = [
    {
        "key": "owasp",
        "icon": "📄",
        "name": "OWASP 标准渗透测试报告",
        "description": "覆盖范围 / 方法论 / 漏洞详情 / 修复建议，业内通用结构",
        "content": """# 渗透测试报告（OWASP 标准）

## 一、报告信息
- 测试目标：
- 测试范围：
- 测试时间：
- 测试人员：
- 报告版本：v1.0

## 二、测试方法
本次测试参照 OWASP Testing Guide v4，覆盖以下类别：
1. 信息收集（Information Gathering）
2. 配置与部署管理测试（Configuration & Deploy Management）
3. 身份认证测试（Authentication Testing）
4. 会话管理测试（Session Management）
5. 输入验证测试（Input Validation）
6. 错误处理（Error Handling）
7. 加密（Cryptography）
8. 业务逻辑（Business Logic）
9. 客户端（Client Side）

## 三、风险等级定义
| 等级 | 定义 |
|------|------|
| 严重 | 可直接获取系统控制权 / 大量数据 |
| 高 | 可越权 / 读敏感数据 |
| 中 | 可被利用，需特定条件 |
| 低 | 信息泄露类 |

## 四、漏洞清单
| 编号 | 漏洞名称 | 风险等级 | URL | 状态 |
|------|----------|----------|-----|------|

## 五、漏洞详情
### 5.1 [漏洞名称]
- **风险等级**：
- **影响范围**：
- **漏洞描述**：
- **复现步骤**：
  1. 
  2. 
- **PoC**：
```
```
- **截图**：
- **修复建议**：

## 六、修复优先级建议
1. **立即修复（72 小时内）**：
2. **限期修复（2 周内）**：
3. **建议修复**：

## 七、总结
"""
    },
    {
        "key": "poc",
        "icon": "🩸",
        "name": "漏洞证明优先（POC-First）报告",
        "description": "突出复现步骤 + 请求/响应包，适合需要客户验证的场景",
        "content": """# 漏洞证明报告

## 摘要
本次测试在目标系统中发现 N 个安全问题，其中严重 X 个、高危 Y 个。本报告以"可复现、可验证"为原则，每个漏洞均提供完整 PoC。

## 漏洞列表
| # | 漏洞 | 等级 | 入口 |
|---|------|------|------|

---

## 漏洞 #1：[名称]

**等级**：严重 / 高 / 中 / 低  
**入口**：`https://...`

### 复现步骤
1. 
2. 
3. 

### 请求包
```http
POST /api/xxx HTTP/1.1
Host: target.com
Content-Type: application/json

{"payload": "..."}
```

### 响应包
```http
HTTP/1.1 200 OK
Content-Type: application/json

{"..."}
```

### 截图
（粘贴截图）

### 影响

### 修复建议

---

## 修复跟进表
| 漏洞 | 负责人 | 计划修复时间 | 状态 |
|------|--------|------------|------|
"""
    },
    {
        "key": "redteam",
        "icon": "🎯",
        "name": "红队对抗报告（攻击链 Timeline）",
        "description": "按时间线展开攻击路径，覆盖 Kill Chain 各阶段",
        "content": """# 红队对抗报告

## 1. 行动概览
- 行动代号：
- 行动时间：YYYY-MM-DD ~ YYYY-MM-DD
- 目标：
- 授权范围：
- 最终战果：（如：拿到域控 / 核心业务库 / 关键数据）

## 2. 攻击链 Timeline
```
T0  踩点  →  T1  外网突破  →  T2  立足点  →  T3  横向  →  T4  权限维持  →  T5  目标达成
```

### 2.1 T0 - 信息收集
- 资产测绘：
- 关键发现：

### 2.2 T1 - 外网突破
- 入口：
- 利用漏洞：
- 时间：

### 2.3 T2 - 立足点
- 落地方式：
- 工具 / 木马：
- 持久化：

### 2.4 T3 - 横向移动
- 路径：A → B → C
- 凭据来源：
- 关键节点：

### 2.5 T4 - 权限维持

### 2.6 T5 - 目标达成

## 3. MITRE ATT&CK 映射
| 战术 | 技术 ID | 描述 |
|------|---------|------|

## 4. 防御检测建议
### 4.1 检出点（Detection）
### 4.2 阻断点（Prevention）
### 4.3 应急响应（Response）

## 5. 复盘
- 防御方做对了什么：
- 防御方欠缺什么：
- 整改优先级：
"""
    },
]


@router.get("/api/templates/builtin")
async def templates_builtin():
    """列出内置示例模版（仅返回元信息列表，不含内容）。"""
    return {
        "status": "ok",
        "templates": [
            {"key": t["key"], "icon": t["icon"], "name": t["name"], "description": t["description"]}
            for t in _BUILTIN_TEMPLATES
        ]
    }


@router.post("/api/templates/builtin/install")
async def templates_builtin_install(request: Request):
    """安装一个内置示例模版到用户模版库。"""
    import uuid as _uuid
    body = await request.json()
    key = (body.get("key") or "").strip()
    src = next((t for t in _BUILTIN_TEMPLATES if t["key"] == key), None)
    if not src:
        return {"status": "error", "message": f"内置模版不存在: {key}"}

    template_id = _uuid.uuid4().hex[:12]
    saved_as = f"{template_id}.md"
    save_path = CUSTOM_TEMPLATE_DIR / saved_as
    content_bytes = src["content"].encode("utf-8")
    save_path.write_bytes(content_bytes)

    text_path = CUSTOM_TEMPLATE_DIR / f"{template_id}.extracted.txt"
    text_path.write_bytes(content_bytes)

    meta = _load_template_meta()
    meta.append({
        "id": template_id,
        "name": src["name"],
        "description": src["description"],
        "filename": f"{src['key']}.md",
        "saved_as": saved_as,
        "ext": ".md",
        "size": len(content_bytes),
        "enabled": False,
        "uploaded_at": time.time(),
        "text_length": len(content_bytes),
        "from_builtin": src["key"],
    })
    _save_template_meta(meta)
    return {"status": "ok", "id": template_id, "name": src["name"]}


@router.post("/api/templates/delete")
async def templates_delete(request: Request):
    """删除自定义模版。"""
    body = await request.json()
    template_id = (body.get("id") or "").strip()

    meta = _load_template_meta()
    new_meta = [m for m in meta if m["id"] != template_id]
    if len(new_meta) == len(meta):
        return {"status": "error", "message": f"模版不存在: {template_id}"}

    for m in meta:
        if m["id"] == template_id:
            file_path = CUSTOM_TEMPLATE_DIR / m["saved_as"]
            if file_path.exists():
                file_path.unlink(missing_ok=True)
            text_path = CUSTOM_TEMPLATE_DIR / f"{template_id}.extracted.txt"
            if text_path.exists():
                text_path.unlink(missing_ok=True)
            break

    _save_template_meta(new_meta)
    return {"status": "ok", "deleted": template_id}


@router.get("/api/templates/preview")
async def templates_preview(template_id: str):
    """预览模版提取的文本内容。"""
    meta = _load_template_meta()
    item = next((m for m in meta if m.get("id") == template_id), None)
    if not item:
        return {"status": "error", "message": "模版不存在"}

    if item.get("ext") == ".doc":
        return {
            "status": "ok",
            "content": (
                "⚠️  .doc（Word 97-2003 老二进制格式）无法被解析。\n\n"
                "建议操作：\n"
                "  1. 用 Word 打开原文件 → 另存为 .docx；\n"
                "  2. 删除该模版后，重新上传 .docx 版本。\n"
            ),
            "unsupported": True,
        }

    text_path = CUSTOM_TEMPLATE_DIR / f"{template_id}.extracted.txt"
    if not text_path.exists():
        return {"status": "error", "message": "模版文本不存在"}
    content = text_path.read_text(encoding="utf-8")
    return {"status": "ok", "content": content[:_TEMPLATE_TEXT_LIMIT]}


@router.post("/api/templates/update")
async def templates_update(request: Request):
    """原地更新现有模版（保持 id / saved_as 不变）。"""
    body = await request.json()
    template_id = (body.get("id") or "").strip()
    name = (body.get("name") or "").strip()
    description = (body.get("description") or "").strip()
    content = body.get("content")

    if not template_id:
        return {"status": "error", "message": "缺少 id"}
    if not name:
        return {"status": "error", "message": "模版名称不能为空"}
    if content is None or not str(content).strip():
        return {"status": "error", "message": "模版内容不能为空"}
    content_bytes = str(content).encode("utf-8")
    if len(content_bytes) > _MAX_TEMPLATE_SIZE:
        return {"status": "error",
                "message": f"内容过大（{len(content_bytes) / 1024 / 1024:.1f}MB），最大 20MB"}

    meta = _load_template_meta()
    item = next((m for m in meta if m.get("id") == template_id), None)
    if not item:
        return {"status": "error", "message": f"模版不存在: {template_id}"}

    ext = item.get("ext", "").lower()
    if ext not in (".md", ".txt", ".html"):
        return {"status": "error",
                "message": f"该格式（{ext}）不支持在线编辑，请删除后重新上传"}

    file_path = CUSTOM_TEMPLATE_DIR / item["saved_as"]
    file_path.write_bytes(content_bytes)

    extracted_text = _extract_template_text(file_path, ext)
    text_path = CUSTOM_TEMPLATE_DIR / f"{template_id}.extracted.txt"
    text_path.write_text(extracted_text, encoding="utf-8")

    item["name"] = name
    item["description"] = description
    item["size"] = len(content_bytes)
    item["text_length"] = len(extracted_text)
    item["updated_at"] = time.time()
    _save_template_meta(meta)

    return {
        "status": "ok",
        "id": template_id,
        "name": name,
        "size": len(content_bytes),
        "text_length": len(extracted_text),
    }


@router.get("/api/templates/active")
async def templates_active():
    """获取当前激活的自定义模版（供报告生成使用）。"""
    meta = _load_template_meta()
    for item in meta:
        if item.get("enabled"):
            text_path = CUSTOM_TEMPLATE_DIR / f"{item['id']}.extracted.txt"
            if text_path.exists():
                return {
                    "status": "ok",
                    "has_custom": True,
                    "template_id": item["id"],
                    "template_name": item["name"],
                    "content": text_path.read_text(encoding="utf-8")[:100000],
                }
    return {"status": "ok", "has_custom": False}
